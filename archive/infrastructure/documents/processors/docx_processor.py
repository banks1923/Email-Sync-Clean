"""DOCX document processor for Microsoft Word files.

Handles DOCX files with text, tables, and metadata extraction.
"""

from pathlib import Path
from typing import Any

from loguru import logger

from .base_processor import BaseProcessor

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class DocxProcessor(BaseProcessor):
    """
    Processes Microsoft Word DOCX documents.
    """

    def __init__(self):
        """
        Initialize DOCX processor.
        """
        super().__init__()
        self.format_type = "docx"
        self.available = DOCX_AVAILABLE

        if not self.available:
            logger.warning("python-docx not installed. Install with: pip install python-docx")

    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        if not self.available:
            raise ImportError("python-docx is required for DOCX processing")

        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            tables_text = self._extract_tables(doc)

            # Extract headers and footers
            headers_footers = self._extract_headers_footers(doc)

            # Combine all text
            all_text = []

            if headers_footers.get("headers"):
                all_text.append("=== HEADERS ===")
                all_text.extend(headers_footers["headers"])
                all_text.append("")

            all_text.extend(paragraphs)

            if tables_text:
                all_text.append("\n=== TABLES ===")
                all_text.extend(tables_text)

            if headers_footers.get("footers"):
                all_text.append("\n=== FOOTERS ===")
                all_text.extend(headers_footers["footers"])

            return "\n".join(all_text)

        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            raise

    def _extract_tables(self, doc: Document) -> list[str]:
        """Extract text from tables in document.

        Args:
            doc: python-docx Document object

        Returns:
            List of table text representations
        """
        tables_text = []

        for i, table in enumerate(doc.tables):
            table_data = []
            table_data.append(f"\nTable {i + 1}:")

            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)

                if any(row_text):  # Skip empty rows
                    table_data.append(" | ".join(row_text))

            if len(table_data) > 1:  # Only add if table has content
                tables_text.append("\n".join(table_data))

        return tables_text

    def _extract_headers_footers(self, doc: Document) -> dict[str, list[str]]:
        """Extract headers and footers from document.

        Args:
            doc: python-docx Document object

        Returns:
            Dictionary with headers and footers text
        """
        result = {"headers": [], "footers": []}

        try:
            # Extract headers
            for section in doc.sections:
                header_text = (
                    section.header.paragraphs[0].text.strip() if section.header.paragraphs else ""
                )
                if header_text and header_text not in result["headers"]:
                    result["headers"].append(header_text)

                footer_text = (
                    section.footer.paragraphs[0].text.strip() if section.footer.paragraphs else ""
                )
                if footer_text and footer_text not in result["footers"]:
                    result["footers"].append(footer_text)

        except Exception as e:
            logger.warning(f"Could not extract headers/footers: {e}")

        return result

    def extract_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract metadata from DOCX file including document properties.

        Args:
            file_path: Path to DOCX file

        Returns:
            Enhanced metadata including document properties
        """
        # Get base metadata
        metadata = super().extract_metadata(file_path)

        if not self.available:
            return metadata

        try:
            doc = Document(file_path)

            # Extract core properties
            core_props = doc.core_properties

            # Add available properties
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.created:
                metadata["doc_created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["doc_modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.revision:
                metadata["revision"] = core_props.revision

            # Document statistics
            stats = self._get_document_stats(doc)
            metadata.update(stats)

        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {e}")

        return metadata

    def _get_document_stats(self, doc: Document) -> dict[str, Any]:
        """Get document statistics.

        Args:
            doc: python-docx Document object

        Returns:
            Document statistics
        """
        stats = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
        }

        # Count styles used
        styles_used = set()
        for para in doc.paragraphs:
            if para.style and para.style.name:
                styles_used.add(para.style.name)

        stats["styles_used"] = list(styles_used)
        stats["unique_styles_count"] = len(styles_used)

        # Check for lists
        list_count = 0
        for para in doc.paragraphs:
            if para.style and ("List" in para.style.name or "Bullet" in para.style.name):
                list_count += 1
        stats["list_items"] = list_count

        # Check for images (relationships)
        try:
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
            stats["image_count"] = image_count
        except Exception:
            stats["image_count"] = 0

        return {"document_stats": stats}

    def extract_comments(self, file_path: Path) -> list[str]:
        """Extract comments from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of comments
        """
        if not self.available:
            return []

        comments = []

        try:
            doc = Document(file_path)

            # Note: python-docx doesn't directly support comment extraction
            # This would require XML parsing for full comment support
            # For now, we get comments from core properties
            if doc.core_properties.comments:
                comments.append(doc.core_properties.comments)

        except Exception as e:
            logger.error(f"Failed to extract comments: {e}")

        return comments

    def validate_content(self, content: str) -> bool:
        """Validate DOCX content.

        Args:
            content: Extracted text

        Returns:
            True if content is valid
        """
        if not super().validate_content(content):
            return False

        # Check for common DOCX extraction issues
        # Sometimes corrupted DOCX files produce repetitive patterns
        lines = content.split("\n")
        if len(lines) > 10:
            # Check for excessive repetition
            unique_lines = set(lines)
            if len(unique_lines) < len(lines) * 0.1:  # Less than 10% unique
                logger.warning("DOCX content appears to have excessive repetition")
                return False

        return True
